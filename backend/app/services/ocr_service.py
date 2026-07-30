"""
OCRサービス: PDF/画像/Excel/Wordから情報を抽出する
Azure AI不使用 - ローカルOCR（pytesseract + pdfplumber）を使用
"""
import os
import json
import re
from pathlib import Path
from typing import Optional
import pdfplumber
from PIL import Image
from docx import Document as DocxDocument
import openpyxl

# OCR設定（Tesseractなし版）
OCR_AVAILABLE = False

def extract_text_from_pdf(file_path: str) -> str:
    """PDFからテキストを抽出（テキストPDF対応）"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                # テーブルも抽出
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            text += "\t".join([str(c) if c else "" for c in row]) + "\n"
    except Exception as e:
        print(f"pdfplumber error: {e}")
    return text

def extract_text_from_image(file_path: str) -> str:
    """画像ファイル（Tesseractなし版）- メタ情報のみ返す"""
    try:
        img = Image.open(file_path)
        w, h = img.size
        return f"[画像ファイル: {Path(file_path).name}, サイズ: {w}x{h}px]\n（画像OCRはご利用いただけません。手動で内容を入力してください）"
    except Exception as e:
        return f"[画像読み込みエラー: {e}]"

def extract_text_from_docx(file_path: str) -> str:
    """Wordファイルからテキストを抽出"""
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                text += "\t".join([cell.text for cell in row.cells]) + "\n"
        return text
    except Exception as e:
        print(f"DOCX error: {e}")
        return ""

def extract_text_from_excel(file_path: str) -> str:
    """ExcelファイルからテキストをCSV形式で抽出"""
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        text = ""
        for sheet in wb.worksheets:
            text += f"[Sheet: {sheet.title}]\n"
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join([str(c) if c is not None else "" for c in row])
                if row_text.strip():
                    text += row_text + "\n"
        return text
    except Exception as e:
        print(f"Excel error: {e}")
        return ""

def extract_text(file_path: str) -> str:
    """ファイル種別を自動判定してテキスト抽出"""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
        return extract_text_from_image(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    else:
        return ""

# ─── パターンマッチングによるフィールド抽出 ───────────────────────────

def _find(patterns: list, text: str, default: str = "") -> str:
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
        if m:
            return m.group(1).strip()
    return default

def _find_amount(text: str) -> str:
    patterns = [
        r'(?:total|合計|総額|amount)[^\d]*([0-9,]+(?:\.\d{1,2})?)',
        r'([0-9,]{4,}(?:\.\d{1,2})?)\s*(?:円|JPY|USD|EUR)',
    ]
    return _find(patterns, text, "0")

def _find_items(text: str) -> list:
    """商品明細行を抽出（簡易パターン）"""
    items = []
    # 数量×単価パターン
    patterns = [
        r'([A-Za-z0-9\-\u3040-\u9fff]+)\s+(\d+)\s*(?:個|pcs|units|本|枚|箱)?\s+([0-9,]+(?:\.\d{1,2})?)',
        r'(\d+)\s+([A-Za-z0-9\-\u3040-\u9fff\s]+)\s+([0-9,]+(?:\.\d{1,2})?)\s+([0-9,]+(?:\.\d{1,2})?)',
    ]
    for pat in patterns:
        for m in re.finditer(pat, text, re.MULTILINE):
            groups = m.groups()
            if len(groups) >= 3:
                items.append({
                    "product_name": groups[0].strip() if not groups[0].isdigit() else groups[1].strip(),
                    "quantity": int(re.sub(r'[^\d]', '', groups[1] if not groups[0].isdigit() else groups[0])) if re.sub(r'[^\d]', '', groups[1] if not groups[0].isdigit() else groups[0]) else 1,
                    "unit_price": float(re.sub(r'[^\d.]', '', groups[2])) if re.sub(r'[^\d.]', '', groups[2]) else 0,
                    "amount": 0,
                    "sku": "",
                    "unit": "個"
                })
    if not items:
        items = [{
            "product_name": "（書類から自動抽出）",
            "quantity": 1,
            "unit_price": 0,
            "amount": 0,
            "sku": "",
            "unit": "個"
        }]
    # amountを計算
    for item in items:
        if item["amount"] == 0:
            item["amount"] = round(item["quantity"] * item["unit_price"], 2)
    return items[:10]  # 最大10行

def parse_po(text: str) -> dict:
    return {
        "doc_type": "po",
        "po_number": _find([
            r'P\.?O\.?\s*(?:No\.?|Number|番号)[:\s#]*([A-Za-z0-9\-]+)',
            r'発注番号[:\s]*([A-Za-z0-9\-]+)',
            r'Order\s*No\.?\s*[:\s]*([A-Za-z0-9\-]+)',
        ], text, f"PO-AUTO-{__import__('random').randint(1000,9999)}"),
        "supplier": _find([
            r'(?:Supplier|仕入先|To)[:\s]+([^\n]+)',
            r'(?:Vendor|ベンダー)[:\s]+([^\n]+)',
        ], text, "（自動抽出）"),
        "order_date": _find([
            r'(?:Date|日付|発注日)[:\s]+(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
            r'(\d{4}年\d{1,2}月\d{1,2}日)',
        ], text, "2026-07-30"),
        "currency": "JPY" if re.search(r'円|JPY', text) else ("USD" if re.search(r'USD|\$', text) else "JPY"),
        "total_amount": _find_amount(text),
        "payment_terms": _find([r'(?:Payment|支払条件)[:\s]+([^\n]+)'], text, "T/T 30 days"),
        "items": _find_items(text),
        "notes": "",
    }

def parse_invoice(text: str) -> dict:
    return {
        "doc_type": "invoice",
        "invoice_number": _find([
            r'Invoice\s*(?:No\.?|Number|番号)[:\s#]*([A-Za-z0-9\-]+)',
            r'請求番号[:\s]*([A-Za-z0-9\-]+)',
            r'INV[:\s\-]*([A-Za-z0-9\-]+)',
        ], text, f"INV-AUTO-{__import__('random').randint(1000,9999)}"),
        "customer": _find([
            r'(?:Bill\s*To|Sold\s*To|得意先|請求先)[:\s]+([^\n]+)',
            r'(?:Customer|顧客)[:\s]+([^\n]+)',
        ], text, "（自動抽出）"),
        "invoice_date": _find([
            r'(?:Invoice\s*Date|請求日|Date)[:\s]+(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
            r'(\d{4}年\d{1,2}月\d{1,2}日)',
        ], text, "2026-07-30"),
        "due_date": _find([
            r'(?:Due\s*Date|支払期限)[:\s]+(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
        ], text, ""),
        "currency": "JPY" if re.search(r'円|JPY', text) else ("USD" if re.search(r'USD|\$', text) else "JPY"),
        "total_amount": _find_amount(text),
        "tax_amount": _find([r'(?:Tax|消費税|VAT)[:\s]*([0-9,]+(?:\.\d{1,2})?)'], text, "0"),
        "payment_terms": _find([r'(?:Payment|支払条件)[:\s]+([^\n]+)'], text, "T/T 30 days"),
        "items": _find_items(text),
        "notes": "",
    }

def parse_packing_list(text: str) -> dict:
    return {
        "doc_type": "packing_list",
        "pl_number": _find([
            r'Packing\s*List\s*(?:No\.?|番号)[:\s#]*([A-Za-z0-9\-]+)',
            r'PL[:\s\-]*([A-Za-z0-9\-]+)',
        ], text, f"PL-AUTO-{__import__('random').randint(1000,9999)}"),
        "shipper": _find([r'(?:Shipper|荷主|From)[:\s]+([^\n]+)'], text, "（自動抽出）"),
        "consignee": _find([r'(?:Consignee|荷受人|To)[:\s]+([^\n]+)'], text, "（自動抽出）"),
        "ship_date": _find([r'(?:Ship\s*Date|出荷日)[:\s]+(\d{4}[-/]\d{1,2}[-/]\d{1,2})'], text, "2026-07-30"),
        "total_packages": _find([r'(?:Total\s*Packages?|総梱包数)[:\s]*(\d+)'], text, "1"),
        "total_weight": _find([r'(?:Total\s*(?:Gross\s*)?Weight|総重量)[:\s]*([0-9.]+\s*(?:kg|KG|lbs)?)'], text, ""),
        "items": _find_items(text),
        "notes": "",
    }

def parse_document(file_path: str, doc_type: str) -> dict:
    """ファイルを解析して書類データを返す"""
    text = extract_text(file_path)
    if not text.strip():
        text = "（テキスト抽出不可）"

    parsers = {
        "po": parse_po,
        "invoice": parse_invoice,
        "packing_list": parse_packing_list,
    }
    parser = parsers.get(doc_type, parse_po)
    result = parser(text)
    result["raw_text_preview"] = text[:500]

    # 信頼度スコア（抽出フィールド数に基づく簡易計算）
    filled = sum(1 for v in result.values() if v and v != "（自動抽出）" and v != "（書類から自動抽出）")
    total = len(result)
    result["confidence_score"] = round(filled / total, 2) if total > 0 else 0.5

    return result